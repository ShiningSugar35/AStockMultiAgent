"""Formal report rendering, atomic publishing, manifests, and recovery."""

from __future__ import annotations

import io
import mimetypes
import os
import shutil
import subprocess
import uuid
from datetime import UTC, datetime
from importlib.metadata import PackageNotFoundError, version
from pathlib import Path
from typing import Any, cast

from astock.core.atomic import atomic_write_bytes
from astock.core.errors import AStockError, FailureClass, PolicyError, StorageError
from astock.core.hashing import canonical_json_bytes, sha256_bytes
from astock.core.object_store import ObjectStore
from astock.core.state import StateStore
from astock.reports.manifest_repository import ReportManifestRepository
from astock.reports.paths import ReportPathResolver, safe_report_file_name
from astock.reports.policy import ReportPolicy, load_report_policy
from astock.reports.preferences import PresentationPreferencesRepository
from astock.reports.validation import validate_docx, validate_pdf, visual_qa_summary
from astock.research.presentation import audit_public_answer
from astock.schemas.presentation import (
    InvestorPresentationModel,
    PublicReportReference,
    ResearchNarrativeBundle,
)
from astock.schemas.reports import (
    AssetManifest,
    AssetRightsStatus,
    CitationLevel,
    PdfConverterCapability,
    PdfPreference,
    PreferenceLength,
    PresentationPreferences,
    PrivacyLevel,
    ReportAsset,
    ReportDirectoryPolicy,
    ReportFormat,
    ReportManifest,
    ReportPublishResult,
    ReportRenderer,
    ReportRequest,
    ReportStatus,
    UnknownRightsAction,
)
from astock.settings import ProjectPaths

_EMBED_RIGHTS = {
    AssetRightsStatus.OWNED,
    AssetRightsStatus.PUBLIC_DOMAIN,
    AssetRightsStatus.LICENSED,
    AssetRightsStatus.PUBLIC_DISCLOSURE,
}


class ReportPublishError(AStockError):
    """Stable report-layer failure."""


class ReportService:
    def __init__(
        self,
        paths: ProjectPaths,
        state: StateStore,
        objects: ObjectStore,
        *,
        policy: ReportPolicy | None = None,
        path_resolver: ReportPathResolver | None = None,
    ) -> None:
        self.paths = paths
        self.state = state
        self.objects = objects
        self.policy = policy or load_report_policy(paths.report_policy)
        self.manifests = ReportManifestRepository(state, objects)
        self.path_resolver = path_resolver or ReportPathResolver(
            controlled_root=paths.reports / "output",
            configured_root=self.policy.directory_policy.server_report_root,
            allow_env_override=self.policy.directory_policy.allow_env_override,
            env_name=self.policy.directory_policy.report_root_env,
        )
        self.preference_defaults = PresentationPreferences(
            default_length=PreferenceLength.STANDARD,
            default_report_format=self.policy.default_format,
            report_directory_policy=ReportDirectoryPolicy.DEFAULT,
            citation_level=self.policy.defaults.citation_level,
            privacy_default=self.policy.defaults.privacy_level,
            pdf_preference=PdfPreference.AUTO,
        )

    def publish(self, request: ReportRequest) -> ReportPublishResult:
        preferences = self._effective_preferences(request.preferences)
        privacy = (
            request.privacy_level
            or preferences.privacy_default
            or self.policy.defaults.privacy_level
        )
        citation_level = (
            request.citation_level
            or preferences.citation_level
            or self.policy.defaults.citation_level
        )
        requested_format = self._requested_format(request, preferences)
        requested_directory_policy = (
            request.directory_policy
            or preferences.report_directory_policy
            or ReportDirectoryPolicy.DEFAULT
        )
        assets, asset_payloads = self._prepare_assets(request)
        input_hashes = self._resolve_artifact_hashes(request)
        request_hash = self._request_hash(
            request,
            assets=assets,
            input_hashes=input_hashes,
            privacy=privacy,
            citation_level=citation_level,
            requested_format=requested_format,
            directory_policy=requested_directory_policy,
            template_version=self.policy.template_version,
        )

        existing = self.manifests.get(request.report_key)
        if existing is not None:
            if existing.request_hash != request_hash:
                return self._conflict_result(
                    request,
                    request_hash,
                    existing,
                    privacy=privacy,
                    citation_level=citation_level,
                    directory_policy=requested_directory_policy,
                )
            recovered = self._recover_terminal(existing)
            if recovered is not None:
                return recovered
            recovered = self._recover_staged(request.report_key)
            if recovered is not None:
                return recovered

        try:
            self.state.register_idempotency(
                f"report:{request.report_key}",
                "report-publish",
                request_hash,
            )
        except ValueError:
            return self._conflict_result(
                request,
                request_hash,
                existing,
                privacy=privacy,
                citation_level=citation_level,
                directory_policy=requested_directory_policy,
            )

        resolved = self.path_resolver.resolve(
            requested_directory_policy,
            custom_root=(
                Path(preferences.custom_report_root)
                if preferences.custom_report_root is not None
                else None
            ),
        )
        safe_base = safe_report_file_name(request.output_name_hint or request.title, max_length=190)
        stage_dir = (self.paths.report_staging / request.report_key).resolve()
        self._ensure_staging_contained(stage_dir)
        if stage_dir.exists():
            shutil.rmtree(stage_dir)
        stage_dir.mkdir(parents=True, exist_ok=True)

        created_at = datetime.now(UTC)
        try:
            text_for_audit = self._render_markdown_text(
                request,
                privacy=privacy,
                citation_level=citation_level,
                assets=assets,
            )
            self._audit_privacy(text_for_audit, privacy)
            (
                stage_path,
                output_format,
                renderer,
                renderer_version,
                converter,
                degradation_reason,
            ) = self._render_chain(
                request,
                stage_dir=stage_dir,
                safe_base=safe_base,
                requested_format=requested_format,
                privacy=privacy,
                citation_level=citation_level,
                assets=assets,
                asset_payloads=asset_payloads,
            )
            self._validate_output(stage_path, output_format)
            output_bytes = stage_path.read_bytes()
            output_hash = sha256_bytes(output_bytes)
            output_name = safe_report_file_name(
                f"{safe_base}-{request.report_key[:12]}.{output_format.value.lower()}"
            )
            final_path = self.path_resolver.final_path(resolved, output_name)
            staged_manifest = ReportManifest(
                report_key=request.report_key,
                request_hash=request_hash,
                input_artifact_ids=request.input_artifact_ids,
                input_artifact_hashes=input_hashes,
                template_version=self.policy.template_version,
                renderer=renderer,
                renderer_version=renderer_version,
                converter=converter,
                output_format=output_format,
                privacy_level=privacy,
                citation_level=citation_level,
                citations=request.citations.model_copy(update={"level": citation_level}),
                assets=assets,
                output_file_name=output_name,
                output_relative_ref=output_name,
                output_sha256=output_hash,
                output_byte_size=len(output_bytes),
                publish_status=ReportStatus.STAGED,
                degradation_reason=degradation_reason,
                publish_attempts=(existing.publish_attempts + 1 if existing else 1),
                destination_policy=resolved.policy,
                created_at=existing.created_at if existing else created_at,
            )
            self.manifests.save(staged_manifest)
            request_object = self.objects.put_json(request.model_dump(mode="json"))
            self.state.set_checkpoint(
                scope_type="report",
                scope_key=request.report_key,
                cursor={
                    "request_object_hash": request_object.sha256,
                    "staged_path": str(stage_path),
                    "final_path": str(final_path),
                    "output_sha256": output_hash,
                    "manifest": staged_manifest.model_dump(mode="json"),
                },
                status=ReportStatus.STAGED.value,
                object_hash=output_hash,
            )
            self._atomic_publish(stage_path, final_path, expected_hash=output_hash)
            return self._finalize(staged_manifest, final_path, recovered=False)
        except Exception as exc:
            self._record_failed(request.report_key, request_hash, existing, exc)
            checkpoint = self.state.get_checkpoint("report", request.report_key)
            cursor = checkpoint.get("cursor") if checkpoint is not None else None
            recoverable_staging = isinstance(cursor, dict) and isinstance(
                cursor.get("manifest"), dict
            )
            if not self.policy.staging.keep_on_error and not recoverable_staging:
                shutil.rmtree(stage_dir, ignore_errors=True)
            raise

    def status(self, report_key: str) -> ReportManifest | None:
        return self.manifests.get(report_key)

    def status_view(self, report_key: str) -> dict[str, object] | None:
        manifest = self.status(report_key)
        if manifest is None:
            return None
        payload: dict[str, object] = manifest.model_dump(mode="json")
        payload["visual_qa"] = None
        checkpoint = self.state.get_checkpoint("report", report_key)
        if (
            checkpoint is not None
            and manifest.output_format is ReportFormat.DOCX
            and isinstance(checkpoint.get("cursor"), dict)
        ):
            final_path = Path(str(checkpoint["cursor"].get("final_path", "")))
            if final_path.is_file():
                try:
                    payload["visual_qa"] = visual_qa_summary(final_path)
                except AStockError:
                    payload["visual_qa"] = {"valid": False}
        return payload

    def recover(self, report_key: str | ReportRequest) -> ReportPublishResult:
        if isinstance(report_key, ReportRequest):
            # Compatibility with the early partial implementation.
            return self.publish(report_key)
        manifest = self.manifests.get(report_key)
        if manifest is None:
            raise ReportPublishError(
                "report recovery state was not found",
                failure_class=FailureClass.STORAGE,
            )
        terminal = self._recover_terminal(manifest)
        if terminal is not None:
            return terminal
        staged = self._recover_staged(report_key)
        if staged is not None:
            return staged
        checkpoint = self.state.get_checkpoint("report", report_key)
        if checkpoint is None or not isinstance(checkpoint.get("cursor"), dict):
            raise ReportPublishError(
                "report recovery checkpoint is unavailable",
                failure_class=FailureClass.STORAGE,
            )
        request_hash = checkpoint["cursor"].get("request_object_hash")
        if not isinstance(request_hash, str) or not self.objects.verify(request_hash):
            raise ReportPublishError(
                "report recovery request is unavailable",
                failure_class=FailureClass.STORAGE,
            )
        request = ReportRequest.model_validate_json(self.objects.get_bytes(request_hash))
        return self.publish(request)

    def _effective_preferences(
        self, request_preferences: PresentationPreferences | None
    ) -> PresentationPreferences:
        persisted = PresentationPreferencesRepository(
            self.state,
            defaults=self.preference_defaults,
        ).get_all()
        if request_preferences is None:
            return persisted
        payload = persisted.model_dump(mode="python")
        for field_name in (
            "default_length",
            "default_report_format",
            "report_directory_policy",
            "custom_report_root",
            "citation_level",
            "privacy_default",
            "pdf_preference",
        ):
            value = getattr(request_preferences, field_name)
            if value is not None:
                payload[field_name] = value
        return PresentationPreferences.model_validate(payload)

    def _resolve_artifact_hashes(self, request: ReportRequest) -> list[str]:
        hashes = list(request.input_artifact_hashes)
        for artifact_id in request.input_artifact_ids:
            row = self.state.artifact_record(artifact_id)
            if row is None:
                raise ReportPublishError(
                    "registered report input is unavailable",
                    failure_class=FailureClass.DATA_QUALITY,
                )
            object_hash = str(row["object_hash"])
            if not self.objects.verify(object_hash):
                raise ReportPublishError(
                    "registered report input failed integrity validation",
                    failure_class=FailureClass.DATA_QUALITY,
                )
            hashes.append(object_hash)
        for value in hashes:
            if not self.objects.verify(value):
                raise ReportPublishError(
                    "report input object is missing or corrupt",
                    failure_class=FailureClass.DATA_QUALITY,
                )
        return list(dict.fromkeys(hashes))

    def _prepare_assets(self, request: ReportRequest) -> tuple[AssetManifest, dict[str, bytes]]:
        if len(request.assets.assets) > self.policy.assets.max_assets:
            raise PolicyError(
                "report asset count exceeds policy",
                failure_class=FailureClass.POLICY_REJECTED,
            )
        prepared: list[ReportAsset] = []
        payloads: dict[str, bytes] = {}
        total = 0
        for asset in request.assets.assets:
            excluded_reason: str | None = asset.exclusion_reason
            if asset.excluded:
                excluded_reason = excluded_reason or "REQUEST_EXCLUDED"
            elif asset.rights not in _EMBED_RIGHTS:
                if self.policy.assets.unknown_rights is UnknownRightsAction.FAIL_CLOSED:
                    raise PolicyError(
                        "asset rights are not sufficient for embedding",
                        failure_class=FailureClass.POLICY_REJECTED,
                    )
                excluded_reason = "RIGHTS_NOT_APPROVED"
            if not request.include_assets:
                excluded_reason = excluded_reason or "ASSET_EMBEDDING_DISABLED"

            raw: bytes | None = None
            digest = asset.object_hash
            byte_size = asset.byte_size
            media_type = asset.media_type
            file_name = asset.file_name
            if excluded_reason is None:
                raw = self._read_asset(asset)
                digest = sha256_bytes(raw)
                if asset.object_hash is not None and asset.object_hash != digest:
                    raise ReportPublishError(
                        "report asset hash mismatch",
                        failure_class=FailureClass.DATA_QUALITY,
                    )
                stored = self.objects.put_bytes(raw)
                digest = stored.sha256
                byte_size = len(raw)
                total += byte_size
                if byte_size > self.policy.assets.max_asset_bytes:
                    raise PolicyError(
                        "report asset exceeds per-asset byte limit",
                        failure_class=FailureClass.POLICY_REJECTED,
                    )
                if total > self.policy.assets.max_total_bytes:
                    raise PolicyError(
                        "report assets exceed total byte limit",
                        failure_class=FailureClass.POLICY_REJECTED,
                    )
                file_name = file_name or self._asset_file_name(asset)
                media_type = media_type or mimetypes.guess_type(file_name)[0]
                if media_type not in self.policy.assets.allowed_media_types:
                    raise PolicyError(
                        "report asset media type is not allowed",
                        failure_class=FailureClass.POLICY_REJECTED,
                    )
                payloads[asset.asset_id] = raw
            sanitized = ReportAsset(
                asset_id=asset.asset_id,
                file_name=safe_report_file_name(file_name) if file_name else None,
                media_type=media_type,
                byte_size=byte_size,
                source_url=asset.source_url,
                source_ref=asset.source_ref,
                object_hash=digest,
                rights=asset.rights,
                rights_note=asset.rights_note,
                alt_text=asset.alt_text,
                caption=asset.caption,
                excluded=excluded_reason is not None,
                exclusion_reason=excluded_reason,
            )
            prepared.append(sanitized)
        return AssetManifest(assets=prepared, total_byte_size=total), payloads

    def _read_asset(self, asset: ReportAsset) -> bytes:
        if asset.object_hash is not None and asset.local_path is None:
            return self.objects.get_bytes(asset.object_hash)
        if asset.local_path is None:
            raise ReportPublishError(
                "embeddable asset has no controlled local object",
                failure_class=FailureClass.DATA_QUALITY,
            )
        path = Path(asset.local_path).expanduser().resolve()
        try:
            return path.read_bytes()
        except OSError as exc:
            raise ReportPublishError(
                "report asset is unavailable",
                failure_class=FailureClass.STORAGE,
            ) from exc

    @staticmethod
    def _asset_file_name(asset: ReportAsset) -> str:
        if asset.local_path:
            return Path(asset.local_path).name
        return f"{asset.asset_id}.bin"

    @staticmethod
    def _request_hash(
        request: ReportRequest,
        *,
        assets: AssetManifest,
        input_hashes: list[str],
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        requested_format: ReportFormat,
        directory_policy: ReportDirectoryPolicy,
        template_version: str,
    ) -> str:
        payload = request.model_dump(
            mode="json",
            exclude={"requested_at", "preferences", "assets", "input_artifact_hashes"},
        )
        payload["assets"] = assets.model_dump(mode="json")
        payload["input_artifact_hashes"] = input_hashes
        payload["privacy_level"] = privacy.value
        payload["citation_level"] = citation_level.value
        payload["requested_format"] = requested_format.value
        payload["directory_policy"] = directory_policy.value
        payload["template_version"] = template_version
        return sha256_bytes(canonical_json_bytes(payload))

    def _requested_format(
        self,
        request: ReportRequest,
        preferences: PresentationPreferences,
    ) -> ReportFormat:
        selected = request.format or preferences.default_report_format or self.policy.default_format
        if request.format is None and preferences.pdf_preference is PdfPreference.PDF_FIRST:
            selected = ReportFormat.PDF
        if (
            selected is ReportFormat.PDF
            and preferences.pdf_preference is PdfPreference.PDF_DISABLED
        ):
            return ReportFormat.MD
        return selected

    def _render_chain(
        self,
        request: ReportRequest,
        *,
        stage_dir: Path,
        safe_base: str,
        requested_format: ReportFormat,
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        assets: AssetManifest,
        asset_payloads: dict[str, bytes],
    ) -> tuple[
        Path,
        ReportFormat,
        ReportRenderer,
        str,
        PdfConverterCapability | None,
        str | None,
    ]:
        if requested_format is ReportFormat.MD:
            path = stage_dir / f"{safe_base}.md"
            atomic_write_bytes(
                path,
                self._render_markdown_text(
                    request,
                    privacy=privacy,
                    citation_level=citation_level,
                    assets=assets,
                ).encode("utf-8"),
            )
            return (
                path,
                ReportFormat.MD,
                ReportRenderer.MARKDOWN,
                "markdown-renderer-v1",
                None,
                None,
            )

        docx_path = stage_dir / f"{safe_base}.docx"
        try:
            docx_bytes = self._render_docx_bytes(
                request,
                privacy=privacy,
                citation_level=citation_level,
                assets=assets,
                asset_payloads=asset_payloads,
            )
            atomic_write_bytes(docx_path, docx_bytes)
            validate_docx(docx_path)
        except Exception as exc:
            if isinstance(exc, (PolicyError, ReportPublishError)):
                raise
            return self._markdown_fallback(
                request,
                stage_dir=stage_dir,
                safe_base=safe_base,
                privacy=privacy,
                citation_level=citation_level,
                assets=assets,
                reason=f"DOCX_RENDER_FAILED:{type(exc).__name__}",
            )

        if requested_format is ReportFormat.DOCX:
            return (
                docx_path,
                ReportFormat.DOCX,
                ReportRenderer.DOCX_PYTHON_DOCX,
                self._docx_renderer_version(),
                None,
                None,
            )

        converter = self._probe_pdf_converter()
        if converter is None:
            return (
                docx_path,
                ReportFormat.DOCX,
                ReportRenderer.DOCX_PYTHON_DOCX,
                self._docx_renderer_version(),
                None,
                "PDF_CONVERTER_UNAVAILABLE",
            )
        pdf_path = stage_dir / f"{safe_base}.pdf"
        try:
            converter_version = self._convert_pdf(converter.converter_id, docx_path, pdf_path)
            validate_pdf(pdf_path)
            capability = converter.model_copy(update={"converter_version": converter_version})
            return (
                pdf_path,
                ReportFormat.PDF,
                ReportRenderer.PDF_EXTERNAL,
                self._docx_renderer_version(),
                capability,
                None,
            )
        except Exception:
            return (
                docx_path,
                ReportFormat.DOCX,
                ReportRenderer.DOCX_PYTHON_DOCX,
                self._docx_renderer_version(),
                converter,
                "PDF_CONVERT_FAILED",
            )

    def _markdown_fallback(
        self,
        request: ReportRequest,
        *,
        stage_dir: Path,
        safe_base: str,
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        assets: AssetManifest,
        reason: str,
    ) -> tuple[
        Path,
        ReportFormat,
        ReportRenderer,
        str,
        PdfConverterCapability | None,
        str | None,
    ]:
        path = stage_dir / f"{safe_base}.md"
        atomic_write_bytes(
            path,
            self._render_markdown_text(
                request,
                privacy=privacy,
                citation_level=citation_level,
                assets=assets,
            ).encode("utf-8"),
        )
        return path, ReportFormat.MD, ReportRenderer.MARKDOWN, "markdown-renderer-v1", None, reason

    def _render_markdown_text(
        self,
        request: ReportRequest,
        *,
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        assets: AssetManifest,
    ) -> str:
        narrative = _coerce_narrative(request)
        lines = [
            f"# {request.title}",
            "",
            f"**主体**：{narrative.subject}",
            "",
            "## 结论",
            narrative.headline,
        ]
        if narrative.conclusion_strength is not None:
            lines.append(f"结论强度：{narrative.conclusion_strength}")
        lines.append("")
        _append_markdown_list(lines, "估值与赔率", narrative.valuation_or_odds)
        _append_markdown_list(lines, "决定性理由", narrative.reasons)
        _append_markdown_list(lines, "主要风险", narrative.risks)
        _append_markdown_list(lines, "改变判断的条件", narrative.change_conditions)
        if narrative.data_as_of is not None:
            lines.extend(["## 数据截至", narrative.data_as_of.isoformat(), ""])
        if citation_level is not CitationLevel.NONE:
            lines.append("## 引用清单")
            for citation in request.citations.citations:
                item = f"- [{citation.citation_id}] {citation.label}"
                if citation.url:
                    item += f" — {citation.url}"
                lines.append(item)
            if not request.citations.citations:
                lines.append("- 本报告未附加可公开展开的引用条目。")
            lines.append("")
        lines.extend(["## 资产与图片权利", f"隐私等级：{privacy.value}"])
        if not assets.assets:
            lines.append("- 未嵌入图片资产。")
        for asset in assets.assets:
            suffix = f"；已排除：{asset.exclusion_reason}" if asset.excluded else "；允许嵌入"
            lines.append(f"- {asset.asset_id}：{asset.rights.value}{suffix}")
        return "\n".join(lines).rstrip() + "\n"

    def _render_docx_bytes(
        self,
        request: ReportRequest,
        *,
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        assets: AssetManifest,
        asset_payloads: dict[str, bytes],
    ) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt

        narrative = _coerce_narrative(request)
        document = Document()
        for style_name, size in (("Normal", 10.5), ("Title", 20), ("Heading 1", 15)):
            style = cast(Any, document.styles[style_name])
            style.font.name = "Times New Roman"
            style.font.size = Pt(size)
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        title = document.add_heading(request.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph(f"主体：{narrative.subject}")
        document.add_paragraph(f"隐私等级：{privacy.value}")
        document.add_heading("结论", level=1)
        document.add_paragraph(narrative.headline)
        if narrative.conclusion_strength is not None:
            document.add_paragraph(f"结论强度：{narrative.conclusion_strength}")
        _docx_list(document, "估值与赔率", narrative.valuation_or_odds)
        _docx_list(document, "决定性理由", narrative.reasons)
        _docx_list(document, "主要风险", narrative.risks)
        _docx_list(document, "改变判断的条件", narrative.change_conditions)
        if narrative.data_as_of is not None:
            document.add_heading("数据截至", level=1)
            document.add_paragraph(narrative.data_as_of.isoformat())
        if citation_level is not CitationLevel.NONE:
            document.add_heading("引用清单", level=1)
            if not request.citations.citations:
                document.add_paragraph("本报告未附加可公开展开的引用条目。")
            for citation in request.citations.citations:
                text = f"[{citation.citation_id}] {citation.label}"
                if citation.url:
                    text += f" — {citation.url}"
                document.add_paragraph(text, style="List Bullet")
        document.add_heading("资产与图片权利", level=1)
        if not assets.assets:
            document.add_paragraph("未嵌入图片资产。")
        for asset in assets.assets:
            suffix = f"；已排除：{asset.exclusion_reason}" if asset.excluded else "；允许嵌入"
            document.add_paragraph(
                f"{asset.asset_id}：{asset.rights.value}{suffix}",
                style="List Bullet",
            )
            raw = asset_payloads.get(asset.asset_id)
            if raw is not None and not asset.excluded:
                document.add_picture(io.BytesIO(raw))
                if asset.caption:
                    paragraph = document.add_paragraph(asset.caption)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        buffer = io.BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _docx_renderer_version() -> str:
        try:
            return f"python-docx-{version('python-docx')}"
        except PackageNotFoundError:
            return "python-docx-unknown"

    def _probe_pdf_converter(self) -> PdfConverterCapability | None:
        if not self.policy.pdf.enabled:
            return None
        converter = self._detect_pdf_converter()
        if converter is None:
            return None
        probed_at = datetime.now(UTC)
        try:
            completed = subprocess.run(
                [converter, "--version"],
                capture_output=True,
                text=True,
                timeout=self.policy.pdf.probe_timeout_seconds,
                check=False,
                env=self._subprocess_environment(self.paths.report_staging),
            )
            if completed.returncode != 0:
                return None
            first_line = (completed.stdout or completed.stderr).strip().splitlines()
            converter_version = first_line[0][:120] if first_line else Path(converter).name
            return PdfConverterCapability(
                probe_id=uuid.uuid4().hex,
                converter_id=converter,
                converter_version=converter_version,
                probed_at=probed_at,
                probe_ok=True,
            )
        except (OSError, subprocess.SubprocessError):
            return None

    def _detect_pdf_converter(self) -> str | None:
        configured = os.environ.get(self.policy.pdf.command_env)
        if configured:
            candidate = Path(configured).expanduser()
            if candidate.is_file():
                return str(candidate.resolve())
            return None
        return shutil.which(self.policy.pdf.command)

    def _convert_pdf(
        self,
        converter: str,
        docx_path: Path,
        pdf_path: Path,
    ) -> str:
        profile = docx_path.parent / "_converter_profile"
        profile.mkdir(parents=True, exist_ok=True)
        command = [
            converter,
            f"-env:UserInstallation={profile.resolve().as_uri()}",
            *self.policy.pdf.cli_args,
            str(pdf_path.parent),
            str(docx_path),
        ]
        completed = subprocess.run(
            command,
            capture_output=True,
            text=True,
            timeout=self.policy.pdf.convert_timeout_seconds,
            check=False,
            env=self._subprocess_environment(docx_path.parent),
        )
        produced = pdf_path.parent / f"{docx_path.stem}.pdf"
        if completed.returncode != 0 or not produced.is_file():
            raise ReportPublishError(
                "PDF converter failed",
                failure_class=FailureClass.CAPABILITY_UNAVAILABLE,
            )
        if produced != pdf_path:
            os.replace(produced, pdf_path)
        return Path(converter).name

    @staticmethod
    def _subprocess_environment(temp_root: Path) -> dict[str, str]:
        environment = dict(os.environ)
        environment["TEMP"] = str(temp_root)
        environment["TMP"] = str(temp_root)
        return environment

    @staticmethod
    def _validate_output(path: Path, output_format: ReportFormat) -> None:
        if output_format is ReportFormat.DOCX:
            validate_docx(path)
        elif output_format is ReportFormat.PDF:
            validate_pdf(path)
        else:
            try:
                text = path.read_text(encoding="utf-8")
            except (OSError, UnicodeError) as exc:
                raise ReportPublishError(
                    "Markdown output failed validation",
                    failure_class=FailureClass.DATA_QUALITY,
                ) from exc
            if not text.strip():
                raise ReportPublishError(
                    "Markdown output is empty",
                    failure_class=FailureClass.DATA_QUALITY,
                )

    @staticmethod
    def _audit_privacy(text: str, privacy: PrivacyLevel) -> None:
        if privacy is not PrivacyLevel.PUBLIC:
            return
        audit = audit_public_answer(text)
        if not audit.safe_to_send:
            raise PolicyError(
                "public report privacy audit failed",
                failure_class=FailureClass.POLICY_REJECTED,
                details={"finding_codes": audit.finding_codes},
            )

    def _ensure_staging_contained(self, stage_dir: Path) -> None:
        try:
            stage_dir.relative_to(self.paths.report_staging.resolve())
        except ValueError as exc:
            raise PolicyError(
                "report staging path escaped the controlled root",
                failure_class=FailureClass.POLICY_REJECTED,
            ) from exc

    @staticmethod
    def _atomic_publish(staged: Path, final: Path, *, expected_hash: str) -> None:
        final.parent.mkdir(parents=True, exist_ok=True)
        if final.exists():
            existing_hash = sha256_bytes(final.read_bytes())
            if existing_hash == expected_hash:
                staged.unlink(missing_ok=True)
                return
            raise PolicyError(
                "a different report already occupies the final path",
                failure_class=FailureClass.CONFLICT,
            )
        temporary = final.parent / f".{final.name}.{uuid.uuid4().hex}.tmp"
        try:
            shutil.copyfile(staged, temporary)
            with temporary.open("r+b") as handle:
                handle.flush()
                os.fsync(handle.fileno())
            if sha256_bytes(temporary.read_bytes()) != expected_hash:
                raise StorageError(
                    "atomic report copy failed integrity verification",
                    failure_class=FailureClass.STORAGE,
                )
            os.replace(temporary, final)
            staged.unlink(missing_ok=True)
        except Exception:
            temporary.unlink(missing_ok=True)
            raise

    def _finalize(
        self,
        staged_manifest: ReportManifest,
        final_path: Path,
        *,
        recovered: bool,
    ) -> ReportPublishResult:
        expected_hash = staged_manifest.output_sha256
        if expected_hash is None or sha256_bytes(final_path.read_bytes()) != expected_hash:
            raise ReportPublishError(
                "published report failed final integrity validation",
                failure_class=FailureClass.STORAGE,
            )
        terminal_status = (
            ReportStatus.DEGRADED if staged_manifest.degradation_reason else ReportStatus.PUBLISHED
        )
        payload = staged_manifest.model_dump(mode="python")
        payload.update(
            {
                "publish_status": terminal_status,
                "recovered_existing": recovered,
                "published_at": datetime.now(UTC),
            }
        )
        terminal = ReportManifest.model_validate(payload)
        self.manifests.save(terminal)
        terminal = self.manifests.finalize_manifest(terminal)
        self.state.set_checkpoint(
            scope_type="report",
            scope_key=terminal.report_key,
            cursor={
                "final_path": str(final_path),
                "output_sha256": terminal.output_sha256,
                "manifest": terminal.model_dump(mode="json"),
            },
            status=terminal.publish_status.value,
            object_hash=terminal.output_sha256,
        )
        return self._result_from_manifest(terminal, recovered=recovered)

    def _recover_terminal(self, manifest: ReportManifest) -> ReportPublishResult | None:
        if manifest.publish_status not in {ReportStatus.PUBLISHED, ReportStatus.DEGRADED}:
            return None
        checkpoint = self.state.get_checkpoint("report", manifest.report_key)
        if checkpoint is None or not isinstance(checkpoint.get("cursor"), dict):
            return None
        final_path_value = checkpoint["cursor"].get("final_path")
        if not isinstance(final_path_value, str) or not manifest.output_sha256:
            return None
        final_path = Path(final_path_value)
        if not final_path.is_file():
            return None
        if sha256_bytes(final_path.read_bytes()) != manifest.output_sha256:
            raise PolicyError(
                "published report hash conflicts with its manifest",
                failure_class=FailureClass.CONFLICT,
            )
        return self._result_from_manifest(manifest, recovered=True)

    def _recover_staged(self, report_key: str) -> ReportPublishResult | None:
        checkpoint = self.state.get_checkpoint("report", report_key)
        if checkpoint is None or not isinstance(checkpoint.get("cursor"), dict):
            return None
        cursor = checkpoint["cursor"]
        manifest_payload = cursor.get("manifest")
        staged_path_value = cursor.get("staged_path")
        final_path_value = cursor.get("final_path")
        if not isinstance(manifest_payload, dict):
            return None
        if not isinstance(staged_path_value, str) or not isinstance(final_path_value, str):
            return None
        staged_path = Path(staged_path_value)
        final_path = Path(final_path_value)
        manifest = ReportManifest.model_validate(manifest_payload)
        if manifest.publish_status is not ReportStatus.STAGED or manifest.output_sha256 is None:
            return None
        output_format = cast(ReportFormat, manifest.output_format)
        if final_path.is_file():
            if sha256_bytes(final_path.read_bytes()) != manifest.output_sha256:
                raise PolicyError(
                    "published report hash conflicts with its staged checkpoint",
                    failure_class=FailureClass.CONFLICT,
                )
            self._validate_output(final_path, output_format)
            staged_path.unlink(missing_ok=True)
            return self._finalize(manifest, final_path, recovered=True)
        if not staged_path.is_file():
            return None
        if sha256_bytes(staged_path.read_bytes()) != manifest.output_sha256:
            raise PolicyError(
                "staged report hash conflicts with its checkpoint",
                failure_class=FailureClass.CONFLICT,
            )
        self._validate_output(staged_path, output_format)
        self._atomic_publish(staged_path, final_path, expected_hash=manifest.output_sha256)
        return self._finalize(manifest, final_path, recovered=True)

    def _record_failed(
        self,
        report_key: str,
        request_hash: str,
        existing: ReportManifest | None,
        exc: Exception,
    ) -> None:
        checkpoint = self.state.get_checkpoint("report", report_key)
        checkpoint_cursor = checkpoint.get("cursor") if checkpoint is not None else None
        if isinstance(checkpoint_cursor, dict) and isinstance(
            checkpoint_cursor.get("manifest"), dict
        ):
            # A fully staged checkpoint is recoverable. Do not destroy its request object,
            # paths, or manifest merely because publication/finalization was interrupted.
            return

        current = self.manifests.get(report_key)
        source = (
            current if current is not None and current.request_hash == request_hash else existing
        )
        failure_class = exc.failure_class if isinstance(exc, AStockError) else FailureClass.INTERNAL
        if source is not None and source.request_hash == request_hash:
            payload = source.model_dump(mode="python")
            payload.update(
                {
                    "publish_status": ReportStatus.FAILED,
                    "degradation_reason": f"{failure_class.value}:{type(exc).__name__}",
                    "published_at": None,
                }
            )
            try:
                failed = ReportManifest.model_validate(payload)
                self.manifests.save(failed)
            except (ValueError, AStockError):
                pass
        self.state.set_checkpoint(
            scope_type="report",
            scope_key=report_key,
            cursor={"request_hash": request_hash, "error_class": type(exc).__name__},
            status=ReportStatus.FAILED.value,
        )

    def _conflict_result(
        self,
        request: ReportRequest,
        request_hash: str,
        existing: ReportManifest | None,
        *,
        privacy: PrivacyLevel,
        citation_level: CitationLevel,
        directory_policy: ReportDirectoryPolicy,
    ) -> ReportPublishResult:
        now = datetime.now(UTC)
        manifest = ReportManifest(
            report_key=request.report_key,
            request_hash=request_hash,
            input_artifact_ids=request.input_artifact_ids,
            input_artifact_hashes=request.input_artifact_hashes,
            template_version=self.policy.template_version,
            renderer=ReportRenderer.MARKDOWN,
            renderer_version="not-rendered",
            privacy_level=privacy,
            citation_level=citation_level,
            citations=request.citations,
            assets=AssetManifest(),
            publish_status=ReportStatus.CONFLICT,
            degradation_reason="REPORT_KEY_CONTENT_CONFLICT",
            publish_attempts=(existing.publish_attempts if existing else 1),
            destination_policy=directory_policy,
            created_at=now,
        )
        return ReportPublishResult(
            report_key=request.report_key,
            publish_status=ReportStatus.CONFLICT,
            degradation_reason=manifest.degradation_reason,
            manifest=manifest,
        )

    @staticmethod
    def _result_from_manifest(
        manifest: ReportManifest,
        *,
        recovered: bool,
    ) -> ReportPublishResult:
        reference = (
            PublicReportReference(file_name=manifest.output_file_name)
            if manifest.output_file_name
            else None
        )
        return ReportPublishResult(
            report_key=manifest.report_key,
            publish_status=manifest.publish_status,
            output_format=manifest.output_format,
            output_sha256=manifest.output_sha256,
            output_file_name=manifest.output_file_name,
            output_relative_ref=manifest.output_relative_ref,
            public_reference=reference,
            degradation_reason=manifest.degradation_reason,
            published_at=manifest.published_at,
            manifest_artifact_id=manifest.manifest_artifact_id,
            recovered_existing=recovered,
            manifest=manifest,
        )


def _coerce_narrative(request: ReportRequest) -> ResearchNarrativeBundle:
    if request.narrative is not None:
        return request.narrative
    if request.presentation is not None:
        presentation: InvestorPresentationModel = request.presentation
        return ResearchNarrativeBundle(
            subject=presentation.subject,
            headline=presentation.headline,
            conclusion_strength=presentation.conclusion_strength,
            valuation_or_odds=presentation.valuation_or_odds,
            reasons=presentation.reasons,
            risks=[presentation.risk] if presentation.risk else [],
            change_conditions=(
                [presentation.change_condition] if presentation.change_condition else []
            ),
            citations=presentation.citations,
        )
    return ResearchNarrativeBundle(subject="研究工件", headline="详见已注册研究工件。")


def _append_markdown_list(lines: list[str], heading: str, values: list[str]) -> None:
    if not values:
        return
    lines.extend([f"## {heading}", *[f"- {value}" for value in values], ""])


def _docx_list(document: object, heading: str, values: list[str]) -> None:
    if not values:
        return
    document.add_heading(heading, level=1)  # type: ignore[attr-defined]
    for value in values:
        document.add_paragraph(value, style="List Bullet")  # type: ignore[attr-defined]


__all__ = ["ReportPublishError", "ReportService"]
